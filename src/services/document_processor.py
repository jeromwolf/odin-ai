"""
문서 처리 모듈 (텍스트 추출 및 마크다운 변환)
"""

import os
import subprocess
import re
import json
import zipfile
import tempfile
from pathlib import Path
from typing import Optional, Dict, Tuple, List
from datetime import datetime
from loguru import logger

# HWP 직접 추출 모듈
try:
    from .hwp_direct_extractor import extract_hwp_safe
    HWP_DIRECT_AVAILABLE = True
except ImportError:
    HWP_DIRECT_AVAILABLE = False
    logger.warning("HWP direct extractor not available")

# 문서 처리 라이브러리
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
    logger.warning("PyPDF2 not installed")

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
    logger.warning("python-docx not installed")

try:
    import pandas as pd
except ImportError:
    pd = None
    logger.warning("pandas not installed")

from sqlalchemy.orm import Session
from src.database.models import BidDocument, BidAnnouncement

# PDF 처리 모듈 import
from src.services.pdf_processor import process_pdf_file

# HWP 고급 처리 모듈 import
from src.services.hwp_advanced_extractor import extract_hwp_advanced

# BeautifulSoup for HTML parsing (optional)
try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None
    logger.warning("BeautifulSoup not installed - advanced table extraction disabled")


class DocumentProcessor:
    """문서 처리기"""

    def __init__(self, db_session: Session, output_path: Path):
        self.db_session = db_session
        self.output_path = Path(output_path)
        self.markdown_path = self.output_path / "markdown"
        self.markdown_path.mkdir(parents=True, exist_ok=True)

        # 통계
        self.stats = {
            'total': 0,
            'success': 0,
            'failed': 0
        }

    def _detect_file_type(self, file_path: Path) -> str:
        """
        파일의 실제 타입을 감지 (확장자가 아닌 파일 내용 기반)

        Returns:
            'hwp', 'hwpx', 'pdf', 'docx', 'xlsx', 'xls', 'unknown'
        """
        try:
            # 파일의 첫 바이트를 읽어 매직 넘버 확인
            with open(file_path, 'rb') as f:
                header = f.read(512)

            # PDF 파일 시그니처: %PDF (가장 먼저 체크 - 명확함)
            if header[:4] == b'%PDF':
                return 'pdf'

            # OLE2 Compound Document (HWP 5.x, XLS, DOC 등)
            if header[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
                # 더 많은 데이터를 읽어서 구분
                with open(file_path, 'rb') as f:
                    content = f.read(8192)  # 8KB 읽기

                    # HWP 5.x 시그니처 우선 확인 (UTF-16LE 인코딩)
                    # OLE2 구조에서 문자열은 대부분 UTF-16LE로 저장됨
                    hwp_signatures = [
                        b'H\x00w\x00p\x00S\x00u\x00m\x00m\x00a\x00r\x00y\x00I\x00n\x00f\x00o\x00r\x00m\x00a\x00t\x00i\x00o\x00n\x00',  # HwpSummaryInformation
                        b'F\x00i\x00l\x00e\x00H\x00e\x00a\x00d\x00e\x00r\x00',  # FileHeader (HWP 특유)
                        b'H\x00W\x00P\x00 \x00D\x00o\x00c\x00u\x00m\x00e\x00n\x00t\x00',  # HWP Document
                        b'H\x00w\x00p\x00D\x00o\x00c\x00',  # HwpDoc
                        b'\xbb\xea\xc7\xd1\xb1\xdb',  # "한글" in EUC-KR
                    ]

                    if any(sig in content for sig in hwp_signatures):
                        return 'hwp'

                    # Excel 시그니처 확인 (UTF-16LE)
                    excel_signatures = [
                        b'M\x00i\x00c\x00r\x00o\x00s\x00o\x00f\x00t\x00 \x00E\x00x\x00c\x00e\x00l\x00',  # Microsoft Excel
                        b'W\x00o\x00r\x00k\x00b\x00o\x00o\x00k\x00',  # Workbook
                        b'Microsoft Excel',  # ASCII 형태도 체크
                        b'Workbook',  # ASCII 형태도 체크
                    ]

                    if any(sig in content for sig in excel_signatures):
                        return 'xls'

                    # Word 시그니처 확인 (UTF-16LE)
                    word_signatures = [
                        b'M\x00i\x00c\x00r\x00o\x00s\x00o\x00f\x00t\x00 \x00W\x00o\x00r\x00d\x00',  # Microsoft Word
                        b'W\x00o\x00r\x00d\x00.\x00D\x00o\x00c\x00u\x00m\x00e\x00n\x00t\x00',  # Word.Document
                        b'Microsoft Word',  # ASCII 형태도 체크
                    ]

                    if any(sig in content for sig in word_signatures):
                        return 'doc'

                    # 위에서 구분 못하면 OLE2이지만 알 수 없음
                    # 파일 확장자 기반으로 추정
                    extension = file_path.suffix.lower()
                    if extension == '.hwp':
                        return 'hwp'
                    elif extension == '.xls':
                        return 'xls'
                    elif extension == '.doc':
                        return 'doc'

            # ZIP 기반 파일들 (HWPX, DOCX, XLSX)
            if header[:4] == b'PK\x03\x04':
                try:
                    with zipfile.ZipFile(file_path, 'r') as z:
                        namelist = z.namelist()

                        # HWPX: Contents/ 디렉토리 포함
                        if any('Contents/' in name or name.startswith('Contents/') for name in namelist):
                            return 'hwpx'

                        # DOCX: word/ 디렉토리 포함
                        if any('word/' in name or name.startswith('word/') for name in namelist):
                            return 'docx'

                        # XLSX: xl/ 디렉토리 포함
                        if any('xl/' in name or name.startswith('xl/') for name in namelist):
                            return 'xlsx'

                        # ZIP이지만 형식 불명 - 확장자로 추정
                        extension = file_path.suffix.lower()
                        if extension == '.hwpx':
                            return 'hwpx'
                        elif extension == '.docx':
                            return 'docx'
                        elif extension == '.xlsx':
                            return 'xlsx'

                except Exception as e:
                    logger.warning(f"ZIP 파일 검사 실패: {e}")
                    # ZIP 파일이지만 열기 실패 - 확장자 기반 추정
                    extension = file_path.suffix.lower()
                    if extension in ['.hwpx', '.docx', '.xlsx']:
                        return extension[1:]

            # 확장자 기반 폴백
            extension = file_path.suffix.lower()
            if extension in ['.hwp', '.hwpx', '.pdf', '.docx', '.xlsx', '.xls', '.doc']:
                return extension[1:]  # . 제거

            return 'unknown'

        except Exception as e:
            logger.error(f"파일 타입 감지 오류: {e}")
            # 에러 발생 시 확장자 기반 폴백
            extension = file_path.suffix.lower()
            if extension:
                return extension[1:]
            return 'unknown'

    async def process_pending_documents(self) -> dict:
        """
        대기 중인 문서 처리
        Returns: 처리 통계
        """
        # 다운로드 완료되었지만 처리되지 않은 문서 조회
        pending_docs = self.db_session.query(BidDocument).filter(
            BidDocument.download_status == 'completed',
            BidDocument.processing_status == 'pending'
        ).all()  # limit 제거 - 모든 pending 문서 처리

        if not pending_docs:
            logger.info("처리할 문서 없음")
            return self.stats

        logger.info(f"{len(pending_docs)}개 문서 처리 시작")

        for doc in pending_docs:
            await self._process_document(doc)

        logger.info(
            f"문서 처리 완료 - 성공: {self.stats['success']}, "
            f"실패: {self.stats['failed']}"
        )

        return self.stats

    async def _process_document(self, document: BidDocument):
        """개별 문서 처리"""
        self.stats['total'] += 1

        try:
            if not document.storage_path:
                document.processing_status = 'failed'
                document.error_message = "파일 경로 없음"
                self.db_session.commit()
                self.stats['failed'] += 1
                return

            file_path = Path(document.storage_path)
            if not file_path.exists():
                document.processing_status = 'failed'
                document.error_message = "파일 없음"
                self.db_session.commit()
                self.stats['failed'] += 1
                return

            # 폴더인 경우 (ZIP 파일이 압축 해제된 경우)
            if file_path.is_dir():
                text_content, extraction_method = await self._extract_hwp(file_path)
            else:
                # 실제 파일 타입 감지 (확장자가 아닌 파일 내용 기반)
                actual_file_type = self._detect_file_type(file_path)
                extension = file_path.suffix.lower()

                # 확장자와 실제 타입이 다른 경우 로그 출력
                expected_type = extension[1:] if extension else 'unknown'
                if actual_file_type != expected_type and actual_file_type != 'unknown':
                    logger.warning(
                        f"⚠️ 파일 타입 불일치 감지: {file_path.name}\n"
                        f"   확장자: {extension} → 실제 타입: {actual_file_type}\n"
                        f"   실제 타입으로 처리합니다."
                    )

                text_content = None
                extraction_method = None

                # 실제 파일 타입에 따라 처리
                if actual_file_type in ['hwp', 'hwpx']:
                    if actual_file_type == 'hwpx':
                        # HWPX 파일은 고급 추출기 사용
                        text_content, extraction_method = await self._extract_hwpx_advanced(file_path)
                    else:
                        # 일반 HWP 파일은 기존 방법 사용
                        text_content, extraction_method = await self._extract_hwp(file_path)
                elif actual_file_type == 'pdf':
                    text_content, extraction_method = await self._extract_pdf(file_path)
                elif actual_file_type in ['docx', 'doc']:
                    text_content, extraction_method = await self._extract_docx(file_path)
                elif actual_file_type in ['xlsx', 'xls']:
                    text_content, extraction_method = await self._extract_excel(file_path)
                else:
                    logger.warning(f"지원하지 않는 형식: {actual_file_type} (파일: {file_path.name})")
                    document.processing_status = 'skipped'
                    document.error_message = f"지원하지 않는 형식: {actual_file_type}"
                    self.db_session.commit()
                    return

            if text_content:
                # 마크다운 변환
                markdown_content = await self._convert_to_markdown(
                    text_content,
                    document
                )

                # 마크다운 파일 저장
                md_file_path = await self._save_markdown(
                    markdown_content,
                    document
                )

                # DB 업데이트
                document.extracted_text = text_content[:10000]  # 처음 10000자만 저장
                document.text_length = len(text_content)
                document.extraction_method = extraction_method
                document.markdown_path = str(md_file_path)
                document.processing_status = 'completed'
                document.processed_at = datetime.now()

                self.db_session.commit()
                self.stats['success'] += 1

                logger.info(
                    f"문서 처리 성공: {file_path.name} "
                    f"({len(text_content):,}자, {extraction_method})"
                )
            else:
                document.processing_status = 'failed'
                document.error_message = "텍스트 추출 실패"
                self.db_session.commit()
                self.stats['failed'] += 1

        except Exception as e:
            document.processing_status = 'failed'
            document.error_message = str(e)
            self.db_session.commit()
            self.stats['failed'] += 1
            logger.error(f"문서 처리 오류: {e}")

    async def _extract_hwp(self, file_path: Path) -> Tuple[Optional[str], str]:
        """HWP 파일 텍스트 추출 (hwp_safe_extractor 사용)"""
        try:
            # hwp_safe_extractor 사용하여 안전하게 처리
            from src.services.hwp_safe_extractor import extract_hwp_safe

            text, method = extract_hwp_safe(file_path)

            if text and not text.startswith("[HWP 문서]"):
                logger.info(f"HWP 처리 성공: {file_path.name} - {len(text)}자 추출 (방법: {method})")
                return text, method
            else:
                # 추출 실패 시 기본 정보만 반환
                file_size = file_path.stat().st_size
                fallback_text = f"""
[HWP 문서 정보]
파일명: {file_path.name}
파일 크기: {file_size:,} bytes
생성일: {datetime.fromtimestamp(file_path.stat().st_ctime).strftime('%Y-%m-%d %H:%M:%S')}

[문서 내용]
(HWP 텍스트 추출 실패. 방법: {method})
"""
                logger.warning(f"HWP 추출 실패, 기본 정보 반환: {file_path.name} (방법: {method})")
                return fallback_text, f"fallback-{method}"

        except Exception as e:
            logger.error(f"HWP 처리 오류: {e}")
            return None, f"error-{str(e)[:50]}"

    async def _extract_hwp_with_tables(self, file_path: Path) -> Optional[str]:
        """HWP 파일에서 표 구조를 포함한 텍스트 추출"""
        # HWP 처리 일시적으로 비활성화
        return None

    def _html_table_to_markdown(self, table_element) -> str:
        """HTML 테이블을 마크다운 테이블로 변환"""
        rows = table_element.find_all('tr')
        if not rows:
            return ""

        md_lines = []
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            cell_texts = [cell.get_text(strip=True).replace('|', '\\|') for cell in cells]
            md_lines.append('| ' + ' | '.join(cell_texts) + ' |')

            # 헤더 구분선 추가
            if i == 0:
                md_lines.append('|' + '---|' * len(cells))

        return '\n'.join(md_lines)

    async def _extract_hwpx_advanced(self, file_path: Path) -> Tuple[Optional[str], str]:
        """HWPX 파일 고급 추출 (XML 파싱, 표, OLE 객체 지원)"""
        try:
            # 고급 HWP 추출기 사용
            result = await extract_hwp_advanced(file_path)

            if result['success'] and result['text']:
                extraction_method = f"hwp-advanced ({len(result['extraction_methods'])}개 방법"
                if result['tables_count'] > 0:
                    extraction_method += f", {result['tables_count']}표"
                if result['ole_objects']:
                    extraction_method += f", {len(result['ole_objects'])}OLE"
                extraction_method += ")"

                logger.info(f"HWPX 고급 처리 완료: {file_path.name} - {len(result['text'])}자, 방법: {result['extraction_methods']}")
                return result['text'], extraction_method
            else:
                # 기본 hwp5txt 백업
                logger.warning(f"HWPX 고급 처리 실패, 기본 방법 시도: {file_path.name}")
                return await self._extract_hwpx_fallback(file_path)

        except Exception as e:
            logger.error(f"HWPX 고급 처리 오류: {e}")
            # 기본 hwp5txt 백업
            return await self._extract_hwpx_fallback(file_path)

    async def _extract_hwpx_fallback(self, file_path: Path) -> Tuple[Optional[str], str]:
        """HWPX 백업 처리 방법 (ZIP XML 직접 파싱)"""
        try:
            import xml.etree.ElementTree as ET

            # HWPX는 ZIP 형식이므로 직접 풀어서 XML 파싱
            with zipfile.ZipFile(file_path, 'r') as zf:
                text_parts = []

                # Contents/section*.xml 파일들에서 텍스트 추출
                section_files = [f for f in zf.namelist()
                               if 'section' in f.lower() and f.endswith('.xml')]

                logger.debug(f"HWPX ZIP 내부 파일: {len(zf.namelist())}개, section 파일: {len(section_files)}개")

                for section_file in section_files[:10]:  # 최대 10개 섹션만 처리
                    try:
                        with zf.open(section_file) as f:
                            content = f.read()

                            # XML 파싱
                            root = ET.fromstring(content)

                            # 모든 텍스트 노드 추출 (재귀적으로)
                            for elem in root.iter():
                                # 태그 이름에 'text', 'para', 'char' 등이 포함된 요소의 텍스트 추출
                                if elem.text and elem.text.strip():
                                    text_parts.append(elem.text.strip())
                                if elem.tail and elem.tail.strip():
                                    text_parts.append(elem.tail.strip())

                    except Exception as e:
                        logger.debug(f"섹션 파일 파싱 실패 ({section_file}): {e}")
                        continue

                # 텍스트 합치기
                if text_parts:
                    full_text = '\n'.join(text_parts)
                    logger.info(f"HWPX ZIP 파싱 성공: {file_path.name} - {len(full_text)}자 추출")
                    return full_text, "hwpx-zip-xml"
                else:
                    logger.warning(f"HWPX ZIP 파싱 실패: 텍스트 없음 ({file_path.name})")
                    return None, "hwpx-zip-no-text"

        except zipfile.BadZipFile:
            logger.error(f"HWPX ZIP 오류: 유효하지 않은 ZIP 파일 ({file_path.name})")
            return None, "hwpx-bad-zip"

        except Exception as e:
            logger.error(f"HWPX 백업 처리 오류: {e}")
            return None, f"hwpx-error-{str(e)[:50]}"

    async def _extract_pdf(self, file_path: Path) -> Tuple[Optional[str], str]:
        """PDF 파일 텍스트 추출 (tools/pdf-viewer 통합)"""
        try:
            # 새로운 PDF 처리기 사용
            result = await process_pdf_file(file_path)

            if result['success'] and result['text']:
                extraction_method = f"tools/pdf-viewer ({result['pages_count']}페이지"
                if result['tables_count'] > 0:
                    extraction_method += f", {result['tables_count']}표"
                extraction_method += ")"

                logger.info(f"PDF 처리 완료: {file_path.name} - {result['pages_count']}페이지, {len(result['text'])}자")
                return result['text'], extraction_method
            else:
                # 기존 PyPDF2 백업 방법
                if PyPDF2:
                    return await self._extract_pdf_fallback(file_path)
                else:
                    logger.error(f"PDF 처리 실패: {file_path.name}")
                    return None, None

        except Exception as e:
            logger.error(f"PDF 처리 오류: {e}")
            # 기존 PyPDF2 백업 방법
            if PyPDF2:
                return await self._extract_pdf_fallback(file_path)
            return None, None

    async def _extract_pdf_fallback(self, file_path: Path) -> Tuple[Optional[str], str]:
        """PDF 백업 처리 방법 (PyPDF2)"""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text_parts = []

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

                if text_parts:
                    logger.info(f"PDF 백업 처리 완료: {file_path.name} - {len(pdf_reader.pages)}페이지")
                    return '\n'.join(text_parts), f'PyPDF2-fallback ({len(pdf_reader.pages)}페이지)'
                else:
                    # 텍스트가 없는 경우 (이미지 기반 PDF)
                    logger.warning(f"PDF 텍스트 없음 (이미지 기반): {file_path.name} - {len(pdf_reader.pages)}페이지")
                    fallback_text = f"""
[PDF 문서 정보]
파일명: {file_path.name}
페이지 수: {len(pdf_reader.pages)}
상태: 이미지 기반 PDF (OCR 필요)

※ 이 PDF는 텍스트를 포함하지 않습니다. OCR (광학 문자 인식) 처리가 필요합니다.
"""
                    return fallback_text, f'PyPDF2-image-based ({len(pdf_reader.pages)}페이지)'

            return None, None

        except Exception as e:
            logger.error(f"PDF 백업 처리 오류: {e}")
            return None, None

    async def _extract_docx(self, file_path: Path) -> Tuple[Optional[str], str]:
        """DOCX 파일 텍스트 추출"""
        if not DocxDocument:
            logger.error("python-docx가 설치되지 않음")
            return None, None

        try:
            doc = DocxDocument(str(file_path))
            text_parts = []

            for para in doc.paragraphs:
                if para.text:
                    text_parts.append(para.text)

            # 표 내용도 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text:
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            if text_parts:
                return '\n'.join(text_parts), 'python-docx'

            return None, None

        except Exception as e:
            logger.error(f"DOCX 처리 오류: {e}")
            return None, None

    async def _extract_excel(self, file_path: Path) -> Tuple[Optional[str], str]:
        """Excel 파일 텍스트 추출"""
        if not pd:
            logger.error("pandas가 설치되지 않음")
            return None, None

        try:
            # 모든 시트 읽기
            excel_file = pd.ExcelFile(file_path)
            text_parts = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                # 시트 이름 추가
                text_parts.append(f"## {sheet_name}")

                # 데이터프레임을 텍스트로 변환
                text_parts.append(df.to_string())

            if text_parts:
                return '\n'.join(text_parts), 'pandas'

            return None, None

        except Exception as e:
            logger.error(f"Excel 처리 오류: {e}")
            return None, None

    async def _convert_to_markdown(
        self,
        text_content: str,
        document: BidDocument
    ) -> str:
        """텍스트를 마크다운으로 변환"""
        # 공고 정보 조회
        announcement = self.db_session.query(BidAnnouncement).filter(
            BidAnnouncement.bid_notice_no == document.bid_notice_no
        ).first()

        # 마크다운 헤더
        markdown = f"# {announcement.title if announcement else document.bid_notice_no}\n\n"

        # 메타데이터
        markdown += "## 📋 공고 정보\n\n"
        if announcement:
            markdown += f"- **공고번호**: {announcement.bid_notice_no}\n"
            markdown += f"- **기관명**: {announcement.organization_name}\n"
            markdown += f"- **공고일**: {announcement.announcement_date}\n"
            markdown += f"- **입찰마감**: {announcement.bid_end_date}\n"

            if announcement.estimated_price:
                markdown += f"- **추정가격**: {announcement.estimated_price:,}원\n"

            if announcement.contract_method:
                markdown += f"- **계약방법**: {announcement.contract_method}\n"

        markdown += f"\n- **문서유형**: {document.document_type}\n"
        markdown += f"- **파일명**: {document.file_name or 'N/A'}\n"
        markdown += f"- **추출일시**: {datetime.now()}\n\n"

        markdown += "---\n\n"

        # 본문 처리
        markdown += "## 📄 문서 내용\n\n"

        # 텍스트 정리 및 구조화
        processed_text = self._process_text_structure(text_content)
        markdown += processed_text

        # 중요 정보 하이라이트
        markdown += "\n\n---\n\n"
        markdown += "## ⚠️ 주요 정보\n\n"

        key_info = self._extract_key_information(text_content)
        for info_type, info_value in key_info.items():
            if info_value:
                markdown += f"- **{info_type}**: {info_value}\n"

        return markdown

    def _process_text_structure(self, text: str) -> str:
        """텍스트 구조화"""
        lines = text.split('\n')
        processed_lines = []
        in_table = False

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 제목 패턴 감지
            if re.match(r'^\d+\.\s+', line):  # 1. 제목
                processed_lines.append(f"\n### {line}")
            elif re.match(r'^[가나다라마바사아자차카타파하]\.\s+', line):  # 가. 제목
                processed_lines.append(f"\n#### {line}")
            elif re.match(r'^\(\d+\)', line):  # (1) 항목
                processed_lines.append(f"- {line}")
            elif '|' in line or '\t' in line:  # 표 데이터
                if not in_table:
                    processed_lines.append("\n```")
                    in_table = True
                processed_lines.append(line)
            else:
                if in_table:
                    processed_lines.append("```\n")
                    in_table = False
                processed_lines.append(line)

        if in_table:
            processed_lines.append("```\n")

        return '\n'.join(processed_lines)

    def _extract_key_information(self, text: str) -> Dict[str, str]:
        """중요 정보 추출"""
        key_info = {}

        # 날짜 패턴
        date_pattern = r'(\d{4}[\.-]\d{1,2}[\.-]\d{1,2})'
        dates = re.findall(date_pattern, text)
        if dates:
            key_info['주요일정'] = ', '.join(set(dates[:5]))  # 처음 5개

        # 금액 패턴
        price_pattern = r'([\d,]+)\s*원'
        prices = re.findall(price_pattern, text)
        if prices:
            # 큰 금액만 표시
            significant_prices = []
            for price_str in prices:
                try:
                    price = int(price_str.replace(',', ''))
                    if price >= 1000000:  # 100만원 이상
                        significant_prices.append(f"{price:,}원")
                except:
                    pass
            if significant_prices:
                key_info['관련금액'] = ', '.join(significant_prices[:3])

        # 자격요건 키워드
        qual_keywords = ['자격', '요건', '등록', '면허', '인증', '허가']
        for keyword in qual_keywords:
            if keyword in text:
                # 키워드 주변 텍스트 추출
                idx = text.find(keyword)
                context = text[max(0, idx-50):min(len(text), idx+100)]
                key_info['자격요건'] = context.strip()[:100] + '...'
                break

        # 제출서류 키워드
        doc_keywords = ['제출서류', '구비서류', '첨부서류', '제출서식']
        for keyword in doc_keywords:
            if keyword in text:
                idx = text.find(keyword)
                context = text[max(0, idx-50):min(len(text), idx+100)]
                key_info['제출서류'] = context.strip()[:100] + '...'
                break

        return key_info

    async def _save_markdown(
        self,
        markdown_content: str,
        document: BidDocument
    ) -> Path:
        """마크다운 파일 저장"""
        # 날짜별 디렉토리
        now = datetime.now()
        date_path = self.markdown_path / f"{now.year}/{now.month:02d}/{now.day:02d}"
        date_path.mkdir(parents=True, exist_ok=True)

        # 파일명 생성
        filename = f"{document.bid_notice_no}_{document.document_type}.md"
        file_path = date_path / filename

        # 파일 저장
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)

        logger.info(f"마크다운 저장: {file_path}")
        return file_path

    def get_processing_stats(self) -> dict:
        """처리 통계 조회"""
        total = self.db_session.query(BidDocument).count()
        completed = self.db_session.query(BidDocument).filter(
            BidDocument.processing_status == 'completed'
        ).count()
        failed = self.db_session.query(BidDocument).filter(
            BidDocument.processing_status == 'failed'
        ).count()
        pending = self.db_session.query(BidDocument).filter(
            BidDocument.processing_status == 'pending'
        ).count()

        return {
            'total': total,
            'completed': completed,
            'failed': failed,
            'pending': pending,
            'completion_rate': f"{(completed/total*100):.1f}%" if total > 0 else "0%"
        }