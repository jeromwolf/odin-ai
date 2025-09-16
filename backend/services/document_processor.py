"""
문서 다운로드 및 처리 서비스 - HWP Viewer 통합 버전
HWP, PDF 등 입찰 관련 문서를 다운로드하고 고도화된 텍스트 추출 및 마크다운 변환
"""

import asyncio
import os
import sys
import hashlib
import tempfile
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from urllib.parse import urlparse
import shutil

import httpx
from loguru import logger

from backend.core.config import settings
from backend.models.database import SessionLocal
from backend.models.document_models import Document, DocumentProcessingQueue

# HWP Viewer 도구들 통합
sys.path.append('/Users/blockmeta/Desktop/blockmeta/project/odin-ai/tools/hwp-viewer')
try:
    from hwp_viewer.parser import HWPParser
    from hwp_viewer.markdown_formatter import MarkdownFormatter
    from hwp_viewer.models import HWPDocument, HWPMetadata
    HWP_VIEWER_AVAILABLE = True
    logger.info("HWP Viewer 도구 로드 완료")
except ImportError as e:
    HWP_VIEWER_AVAILABLE = False
    logger.warning(f"HWP Viewer 도구 로드 실패: {e} - hwp5txt 명령어로 대체")


class DocumentProcessor:
    """문서 다운로드 및 처리 서비스"""

    def __init__(self):
        """프로세서 초기화"""
        self.base_path = Path("storage")
        self.downloads_path = self.base_path / "downloads"
        self.processed_path = self.base_path / "processed"
        self.markdown_path = self.base_path / "markdown"

        # 디렉토리 생성
        self._ensure_directories()

        # HWP Viewer 도구 초기화
        if HWP_VIEWER_AVAILABLE:
            self.hwp_parser = HWPParser()
            self.markdown_formatter = MarkdownFormatter(use_emoji=True)
            logger.info("HWP Viewer 파서 및 포맷터 초기화 완료")
        else:
            self.hwp_parser = None
            self.markdown_formatter = None

        # 지원 파일 타입
        self.supported_types = {
            ".hwp": "hwp",
            ".pdf": "pdf",
            ".doc": "doc",
            ".docx": "doc",
            ".xls": "excel",
            ".xlsx": "excel",
            ".zip": "archive",
            ".rar": "archive"
        }

        # 처리 통계
        self.stats = {
            "downloaded": 0,
            "processed": 0,
            "failed": 0,
            "skipped": 0,
            "hwp_viewer_processed": 0,
            "hwp5txt_processed": 0
        }

        logger.info("DocumentProcessor 초기화 완료 - HWP Viewer 통합 버전")

    def _ensure_directories(self):
        """필요한 디렉토리 생성"""
        dirs_to_create = [
            self.downloads_path / "hwp",
            self.downloads_path / "pdf",
            self.downloads_path / "doc",
            self.downloads_path / "archive",
            self.downloads_path / "unknown",
            self.processed_path / "hwp",
            self.processed_path / "pdf",
            self.processed_path / "doc",
            self.processed_path / "archive",
            self.processed_path / "unknown",
            self.markdown_path
        ]

        for dir_path in dirs_to_create:
            dir_path.mkdir(parents=True, exist_ok=True)

        logger.debug(f"디렉토리 구조 생성 완료: {self.base_path}")

    def get_file_type(self, filename: str) -> str:
        """파일 확장자로 타입 결정"""
        if not filename:
            return "unknown"

        ext = Path(filename).suffix.lower()
        return self.supported_types.get(ext, "unknown")

    def generate_file_hash(self, file_path: Path) -> str:
        """파일 해시 생성 (중복 방지용)"""
        try:
            hasher = hashlib.sha256()
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hasher.update(chunk)
            return hasher.hexdigest()
        except Exception as e:
            logger.error(f"해시 생성 실패: {e}")
            return ""

    async def download_document(
        self,
        url: str,
        filename: str = None,
        bid_notice_no: str = None
    ) -> Dict[str, Any]:
        """문서 다운로드"""

        if not filename:
            # URL에서 파일명 추출
            parsed_url = urlparse(url)
            filename = Path(parsed_url.path).name
            if not filename:
                filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        file_type = self.get_file_type(filename)
        download_dir = self.downloads_path / file_type

        logger.info(f"문서 다운로드 시작: {filename} (타입: {file_type})")

        try:
            # 안전한 파일명 생성
            safe_filename = self._make_safe_filename(filename, bid_notice_no)
            file_path = download_dir / safe_filename

            # 기존 파일 확인
            if file_path.exists():
                logger.info(f"파일 이미 존재: {safe_filename}")
                return await self._create_download_result(file_path, url, "exists")

            # HTTP 다운로드
            async with httpx.AsyncClient(timeout=60.0, follow_redirects=True) as client:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36',
                    'Accept': 'application/octet-stream, application/*, */*'
                }

                response = await client.get(url, headers=headers)
                response.raise_for_status()

                # 파일 저장
                with open(file_path, 'wb') as f:
                    f.write(response.content)

                file_size = file_path.stat().st_size
                logger.info(f"다운로드 완료: {safe_filename} ({file_size:,} bytes)")

                self.stats["downloaded"] += 1
                return await self._create_download_result(file_path, url, "success")

        except httpx.HTTPStatusError as e:
            logger.error(f"HTTP 오류 [{e.response.status_code}]: {url}")
            self.stats["failed"] += 1
            return {
                "success": False,
                "error": f"HTTP {e.response.status_code}",
                "url": url,
                "filename": filename
            }
        except Exception as e:
            logger.error(f"다운로드 실패: {e}")
            self.stats["failed"] += 1
            return {
                "success": False,
                "error": str(e),
                "url": url,
                "filename": filename
            }

    def _make_safe_filename(self, filename: str, bid_notice_no: str = None) -> str:
        """안전한 파일명 생성"""
        # 위험한 문자 제거
        safe_chars = "-_.() abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789가-힣"
        safe_name = ''.join(c for c in filename if c in safe_chars)

        # 길이 제한
        name_part = Path(safe_name).stem[:100]
        ext_part = Path(safe_name).suffix

        # 입찰공고번호 추가
        if bid_notice_no:
            prefix = f"{bid_notice_no}_"
        else:
            prefix = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_"

        return f"{prefix}{name_part}{ext_part}"

    async def _create_download_result(self, file_path: Path, url: str, status: str) -> Dict[str, Any]:
        """다운로드 결과 생성"""
        file_hash = self.generate_file_hash(file_path)
        file_size = file_path.stat().st_size
        file_type = self.get_file_type(file_path.name)

        return {
            "success": True,
            "status": status,
            "file_path": str(file_path),
            "filename": file_path.name,
            "file_type": file_type,
            "file_size": file_size,
            "file_hash": file_hash,
            "url": url,
            "downloaded_at": datetime.now().isoformat()
        }

    async def process_document(self, file_path: str) -> Dict[str, Any]:
        """문서 텍스트 추출 및 마크다운 생성"""

        file_path = Path(file_path)
        if not file_path.exists():
            return {
                "success": False,
                "error": f"파일이 존재하지 않음: {file_path}"
            }

        file_type = self.get_file_type(file_path.name)
        logger.info(f"문서 처리 시작: {file_path.name} (타입: {file_type})")

        try:
            if file_type == "hwp":
                return await self._process_hwp(file_path)
            elif file_type == "pdf":
                return await self._process_pdf(file_path)
            elif file_type == "doc":
                return await self._process_doc(file_path)
            elif file_type == "archive":
                return await self._process_archive(file_path)
            else:
                logger.warning(f"지원하지 않는 파일 타입: {file_type}")
                self.stats["skipped"] += 1
                return {
                    "success": False,
                    "error": f"지원하지 않는 파일 타입: {file_type}",
                    "file_type": file_type
                }

        except Exception as e:
            logger.error(f"문서 처리 실패: {e}")
            self.stats["failed"] += 1
            return {
                "success": False,
                "error": str(e),
                "file_path": str(file_path)
            }

    async def _process_hwp(self, file_path: Path) -> Dict[str, Any]:
        """HWP 파일 처리 - HWP Viewer 통합 버전"""
        processed_dir = self.processed_path / "hwp"
        base_name = file_path.stem

        # 출력 파일 경로
        txt_file = processed_dir / f"{base_name}.txt"
        md_file = self.markdown_path / f"{base_name}.md"

        # HWP Viewer 도구 우선 사용
        if HWP_VIEWER_AVAILABLE and self.hwp_parser and self.markdown_formatter:
            try:
                logger.info(f"HWP Viewer로 파일 처리 시작: {file_path.name}")

                # HWP 문서 파싱
                hwp_document = self.hwp_parser.parse(str(file_path))
                extracted_text = hwp_document.raw_text

                if not extracted_text or len(extracted_text.strip()) < 10:
                    logger.warning("HWP Viewer에서 충분한 텍스트를 추출하지 못함, hwp5txt로 대체")
                    return await self._process_hwp_fallback(file_path)

                # 메타데이터 생성
                metadata = {
                    "파일명": file_path.name,
                    "파일크기": f"{file_path.stat().st_size:,} bytes",
                    "처리일시": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "처리방법": "HWP Viewer 고도화 파서",
                    "단락수": len(hwp_document.paragraphs),
                    "테이블수": len(hwp_document.tables) if hwp_document.tables else 0
                }

                # 문서 메타데이터 추가
                if hwp_document.metadata:
                    if hwp_document.metadata.title:
                        metadata["원본제목"] = hwp_document.metadata.title
                    if hwp_document.metadata.author:
                        metadata["작성자"] = hwp_document.metadata.author
                    if hwp_document.metadata.created_date:
                        metadata["작성일"] = str(hwp_document.metadata.created_date)

                # 텍스트 파일 저장
                with open(txt_file, 'w', encoding='utf-8') as f:
                    f.write(extracted_text)

                # 고도화된 마크다운 생성
                title = hwp_document.metadata.title if hwp_document.metadata and hwp_document.metadata.title else base_name
                markdown_content = self.markdown_formatter.format_document(
                    title=title,
                    content=extracted_text,
                    metadata=metadata
                )

                with open(md_file, 'w', encoding='utf-8') as f:
                    f.write(markdown_content)

                self.stats["processed"] += 1
                self.stats["hwp_viewer_processed"] += 1

                logger.info(f"HWP Viewer 처리 완료: {len(extracted_text)} 문자 추출")

                return {
                    "success": True,
                    "file_type": "hwp",
                    "processing_method": "hwp_viewer",
                    "original_file": str(file_path),
                    "text_file": str(txt_file),
                    "markdown_file": str(md_file),
                    "text_length": len(extracted_text),
                    "paragraph_count": len(hwp_document.paragraphs),
                    "table_count": len(hwp_document.tables) if hwp_document.tables else 0,
                    "metadata": metadata,
                    "processed_at": datetime.now().isoformat()
                }

            except Exception as e:
                logger.error(f"HWP Viewer 처리 실패: {e}, hwp5txt로 대체")
                return await self._process_hwp_fallback(file_path)
        else:
            # HWP Viewer 사용 불가시 hwp5txt 사용
            return await self._process_hwp_fallback(file_path)

    async def _process_hwp_fallback(self, file_path: Path) -> Dict[str, Any]:
        """HWP 파일 처리 - hwp5txt 대체 방법"""
        processed_dir = self.processed_path / "hwp"
        base_name = file_path.stem

        txt_file = processed_dir / f"{base_name}.txt"
        md_file = self.markdown_path / f"{base_name}.md"

        try:
            # hwp5txt 명령으로 텍스트 추출
            result = subprocess.run([
                'hwp5txt', str(file_path)
            ], capture_output=True, text=True, encoding='utf-8')

            if result.returncode != 0:
                logger.warning(f"hwp5txt 실패, 오류: {result.stderr}")
                return {
                    "success": False,
                    "error": f"hwp5txt 처리 실패: {result.stderr}"
                }

            extracted_text = result.stdout

            # 텍스트 파일 저장
            with open(txt_file, 'w', encoding='utf-8') as f:
                f.write(extracted_text)

            # 기본 마크다운 파일 생성
            markdown_content = self._convert_text_to_markdown(extracted_text, file_path.name)
            with open(md_file, 'w', encoding='utf-8') as f:
                f.write(markdown_content)

            self.stats["processed"] += 1
            self.stats["hwp5txt_processed"] += 1

            return {
                "success": True,
                "file_type": "hwp",
                "processing_method": "hwp5txt",
                "original_file": str(file_path),
                "text_file": str(txt_file),
                "markdown_file": str(md_file),
                "text_length": len(extracted_text),
                "processed_at": datetime.now().isoformat()
            }

        except FileNotFoundError:
            logger.error("hwp5txt 명령을 찾을 수 없습니다. HWP 도구가 설치되지 않음")
            return {
                "success": False,
                "error": "hwp5txt 명령을 찾을 수 없음 (HWP 도구 미설치)"
            }

    async def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """PDF 파일 처리"""
        processed_dir = self.processed_path / "pdf"
        base_name = file_path.stem

        # 출력 파일 경로
        txt_file = processed_dir / f"{base_name}.txt"
        md_file = processed_dir / f"{base_name}.md"

        try:
            # pdfplumber로 텍스트 추출 (별도 패키지 필요)
            extracted_text = await self._extract_pdf_text(file_path)

            if not extracted_text:
                return {
                    "success": False,
                    "error": "PDF에서 텍스트를 추출할 수 없음"
                }

            # 텍스트 파일 저장
            with open(txt_file, 'w', encoding='utf-8') as f:
                f.write(extracted_text)

            # 마크다운 파일 생성
            markdown_content = self._convert_text_to_markdown(extracted_text, file_path.name)
            with open(md_file, 'w', encoding='utf-8') as f:
                f.write(markdown_content)

            self.stats["processed"] += 1

            return {
                "success": True,
                "file_type": "pdf",
                "original_file": str(file_path),
                "text_file": str(txt_file),
                "markdown_file": str(md_file),
                "text_length": len(extracted_text),
                "processed_at": datetime.now().isoformat()
            }

        except Exception as e:
            logger.error(f"PDF 처리 실패: {e}")
            return {
                "success": False,
                "error": f"PDF 처리 실패: {e}"
            }

    async def _extract_pdf_text(self, file_path: Path) -> str:
        """PDF에서 텍스트 추출 (pdfplumber 사용)"""
        try:
            import pdfplumber

            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_content.append(f"\n--- 페이지 {page_num} ---\n")
                        text_content.append(text)

            return '\n'.join(text_content)

        except ImportError:
            logger.error("pdfplumber가 설치되지 않음")
            return ""
        except Exception as e:
            logger.error(f"PDF 텍스트 추출 실패: {e}")
            return ""

    async def _process_doc(self, file_path: Path) -> Dict[str, Any]:
        """DOC/DOCX 파일 처리"""
        processed_dir = self.processed_path / "doc"
        base_name = file_path.stem

        # 출력 파일 경로
        txt_file = processed_dir / f"{base_name}.txt"
        md_file = processed_dir / f"{base_name}.md"

        try:
            # python-docx로 텍스트 추출 (별도 패키지 필요)
            extracted_text = await self._extract_doc_text(file_path)

            if not extracted_text:
                return {
                    "success": False,
                    "error": "DOC에서 텍스트를 추출할 수 없음"
                }

            # 텍스트 파일 저장
            with open(txt_file, 'w', encoding='utf-8') as f:
                f.write(extracted_text)

            # 마크다운 파일 생성
            markdown_content = self._convert_text_to_markdown(extracted_text, file_path.name)
            with open(md_file, 'w', encoding='utf-8') as f:
                f.write(markdown_content)

            self.stats["processed"] += 1

            return {
                "success": True,
                "file_type": "doc",
                "original_file": str(file_path),
                "text_file": str(txt_file),
                "markdown_file": str(md_file),
                "text_length": len(extracted_text),
                "processed_at": datetime.now().isoformat()
            }

        except Exception as e:
            logger.error(f"DOC 처리 실패: {e}")
            return {
                "success": False,
                "error": f"DOC 처리 실패: {e}"
            }

    async def _extract_doc_text(self, file_path: Path) -> str:
        """DOC/DOCX에서 텍스트 추출"""
        try:
            from docx import Document

            doc = Document(file_path)
            text_content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            return '\n'.join(text_content)

        except ImportError:
            logger.error("python-docx가 설치되지 않음")
            return ""
        except Exception as e:
            logger.error(f"DOC 텍스트 추출 실패: {e}")
            return ""

    async def _process_archive(self, file_path: Path) -> Dict[str, Any]:
        """압축 파일 처리 (ZIP, RAR)"""
        processed_dir = self.processed_path / "archive"
        base_name = file_path.stem

        try:
            logger.info(f"압축 파일 처리 시작: {file_path.name}")

            # 임시 디렉토리에 압축 해제
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_path = Path(temp_dir)

                # 압축 해제
                if file_path.suffix.lower() == '.zip':
                    import zipfile
                    try:
                        with zipfile.ZipFile(file_path, 'r') as zip_ref:
                            zip_ref.extractall(temp_path)
                    except zipfile.BadZipFile:
                        return {
                            "success": False,
                            "error": "손상된 ZIP 파일"
                        }
                elif file_path.suffix.lower() == '.rar':
                    # RAR 처리는 별도 라이브러리 필요
                    logger.warning("RAR 파일 처리는 아직 지원되지 않음")
                    return {
                        "success": False,
                        "error": "RAR 파일 처리 미지원"
                    }

                # 해제된 파일 목록 확인
                extracted_files = list(temp_path.rglob('*'))
                processable_files = [
                    f for f in extracted_files
                    if f.is_file() and f.suffix.lower() in ['.hwp', '.pdf', '.docx', '.doc', '.txt']
                ]

                if not processable_files:
                    return {
                        "success": False,
                        "error": "압축 파일 내에 처리 가능한 문서가 없음"
                    }

                # 첫 번째 처리 가능한 파일을 메인으로 처리
                main_file = processable_files[0]
                logger.info(f"압축 파일 내 메인 문서 처리: {main_file.name}")

                # 임시 파일을 다운로드 디렉토리로 복사
                temp_target = self.downloads_path / self.get_file_type(main_file.name) / main_file.name
                temp_target.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(main_file, temp_target)

                # 재귀적으로 파일 처리
                main_result = await self.process_document(str(temp_target))

                # 임시 파일 정리
                if temp_target.exists():
                    temp_target.unlink()

                # 압축 파일 정보와 처리 결과 결합
                archive_result = {
                    "success": main_result.get("success", False),
                    "file_type": "archive",
                    "processing_method": "archive_extraction",
                    "original_file": str(file_path),
                    "archive_contents": [f.name for f in processable_files],
                    "processed_file": main_file.name,
                    "extracted_count": len(processable_files),
                    "processed_at": datetime.now().isoformat()
                }

                # 메인 파일 처리 결과 병합
                if main_result.get("success"):
                    archive_result.update({
                        "text_file": main_result.get("text_file"),
                        "markdown_file": main_result.get("markdown_file"),
                        "text_length": main_result.get("text_length", 0),
                        "metadata": main_result.get("metadata", {})
                    })

                    self.stats["processed"] += 1
                else:
                    archive_result["error"] = main_result.get("error", "처리 실패")

                return archive_result

        except Exception as e:
            logger.error(f"압축 파일 처리 실패: {e}")
            return {
                "success": False,
                "error": f"압축 파일 처리 실패: {str(e)}",
                "file_type": "archive"
            }

    def _convert_text_to_markdown(self, text: str, filename: str) -> str:
        """텍스트를 마크다운 형식으로 변환"""
        lines = text.split('\n')
        markdown_lines = [
            f"# {Path(filename).stem}",
            "",
            f"**파일명**: {filename}",
            f"**처리일시**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"**텍스트 길이**: {len(text):,} 문자",
            "",
            "---",
            ""
        ]

        current_section = ""
        for line in lines:
            line = line.strip()
            if not line:
                markdown_lines.append("")
                continue

            # 섹션 헤더 감지 (간단한 패턴)
            if any(keyword in line for keyword in ["1.", "가.", "○", "■", "◦"]):
                if len(line) < 100:  # 헤더는 보통 짧음
                    markdown_lines.append(f"## {line}")
                    current_section = line
                    continue

            # 중요 정보 강조
            if any(keyword in line for keyword in ["예정가격", "추정가격", "입찰마감", "개찰일시", "계약기간"]):
                markdown_lines.append(f"**{line}**")
            else:
                markdown_lines.append(line)

        return '\n'.join(markdown_lines)

    async def download_and_process_batch(
        self,
        document_urls: List[Dict[str, str]],
        bid_notice_no: str = None
    ) -> Dict[str, Any]:
        """문서 일괄 다운로드 및 처리"""

        logger.info(f"일괄 처리 시작: {len(document_urls)}개 문서")

        results = {
            "total_documents": len(document_urls),
            "downloaded": 0,
            "processed": 0,
            "failed": 0,
            "results": [],
            "stats": self.stats.copy()
        }

        for doc_info in document_urls:
            url = doc_info.get("url", "")
            filename = doc_info.get("filename", "")

            if not url:
                continue

            try:
                # 1단계: 다운로드
                download_result = await self.download_document(url, filename, bid_notice_no)

                if download_result.get("success"):
                    results["downloaded"] += 1

                    # 2단계: 처리 (텍스트 추출 + 마크다운 생성)
                    file_path = download_result["file_path"]
                    process_result = await self.process_document(file_path)

                    if process_result.get("success"):
                        results["processed"] += 1

                    # 결과 병합
                    combined_result = {**download_result, **process_result}
                    results["results"].append(combined_result)

                else:
                    results["failed"] += 1
                    results["results"].append(download_result)

            except Exception as e:
                logger.error(f"문서 처리 오류 ({url}): {e}")
                results["failed"] += 1
                results["results"].append({
                    "success": False,
                    "error": str(e),
                    "url": url,
                    "filename": filename
                })

        # 최종 통계 업데이트
        results["final_stats"] = self.stats.copy()

        logger.info(f"일괄 처리 완료: {results['processed']}/{results['total_documents']}개 처리 성공")

        return results

    async def get_processing_stats(self) -> Dict[str, Any]:
        """처리 통계 조회"""
        return {
            "stats": self.stats,
            "directories": {
                "downloads": str(self.downloads_path),
                "processed": str(self.processed_path)
            },
            "supported_types": list(self.supported_types.keys()),
            "timestamp": datetime.now().isoformat()
        }

    async def cleanup_old_files(self, days_old: int = 30) -> Dict[str, Any]:
        """오래된 파일 정리"""
        cutoff_time = datetime.now().timestamp() - (days_old * 24 * 3600)

        removed_files = {"downloads": 0, "processed": 0}

        for category in ["downloads", "processed"]:
            base_dir = self.downloads_path if category == "downloads" else self.processed_path

            for file_path in base_dir.rglob("*"):
                if file_path.is_file() and file_path.stat().st_mtime < cutoff_time:
                    try:
                        file_path.unlink()
                        removed_files[category] += 1
                    except Exception as e:
                        logger.warning(f"파일 삭제 실패: {e}")

        logger.info(f"파일 정리 완료: {sum(removed_files.values())}개 파일 삭제")

        return {
            "removed_files": removed_files,
            "days_old": days_old,
            "cleaned_at": datetime.now().isoformat()
        }